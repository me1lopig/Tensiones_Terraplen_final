import openpyxl
from docx import Document
from docx.shared import Cm, Pt

# Ruta del archivo Excel
excel_file = 'datos_terreno.xlsx'

# Abre el archivo Excel
workbook = openpyxl.load_workbook(excel_file)

# Selecciona la hoja de Excel
sheet = workbook['Hoja1']

# Crea un nuevo documento de Word
doc = Document()

# Crea una tabla en el documento de Word
table = doc.add_table(rows=0, cols=sheet.max_column)
table.autofit = False
# Establecer el tamaño de fuente para todas las celdas de la tabla
style = doc.styles['Normal']
font = style.font
font.size = Pt(9)


# Obtén los datos de Excel y añádelos a la tabla de Word
for row in sheet.iter_rows():
    table_row = table.add_row().cells
    for i, cell in enumerate(row):
        table_row[i].text = str(cell.value)


# Ajusta el ancho de las celdas para que se adecuen al contenido
for column in table.columns:
    for cell in column.cells:
        cell.width=Cm(1.5)


text = "\n\n"
paragraph = doc.add_paragraph(text)

# Guarda el documento de Word
doc.save('ruta_del_archivo.docx')